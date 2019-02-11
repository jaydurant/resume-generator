from . import Resume
from docx import Document
from docx.shared import *

class SerifResumeBuilder(Resume):
    NAME_FONT_SIZE = 14

    def build(self):
        self._document = Document()
        document = self._document

        #Name
        name = self._build_name(self._first_name, self._middle_name, self._last_name)
        p_name = document.add_paragraph(text=name, style='Title')
        p_name_font = p_name.add_run().font
        p_name_font.size = Pt(14)
        p_name_font.bold = True
    
        #

